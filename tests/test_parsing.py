from __future__ import annotations

import base64
from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook

from freca.manifest import sha256_file
from freca.models import ContentKind, SourceRecord, SourceType
from freca.parsing.docx import parse_docx
from freca.parsing.pdf import parse_pdf
from freca.parsing.xlsx import parse_xlsx


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _source(path: Path, source_type: SourceType, track: int = 2) -> SourceRecord:
    return SourceRecord(
        source_id=f"case-001-t{track}",
        case_id=1,
        track=track,
        re_number="RE-WA-2021-0041",
        path=path,
        source_type=source_type,
        sha256=sha256_file(path),
    )


def test_docx_parser_preserves_paragraph_table_and_image_provenance(
    tmp_path: Path,
) -> None:
    image = tmp_path / "pixel.png"
    image.write_bytes(_ONE_PIXEL_PNG)
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("HACCP Plan", level=1)
    doc.add_paragraph("The registered establishment maintains a plan.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Bait stations"
    table.cell(1, 1).text = "Inspected"
    doc.add_picture(str(image))
    doc.save(path)

    chunks = parse_docx(_source(path, SourceType.DOCX), tmp_path / "images")

    assert any(c.content_kind == ContentKind.HEADING and c.content == "HACCP Plan" for c in chunks)
    assert any(c.content_kind == ContentKind.PARAGRAPH and "maintains" in c.content for c in chunks)
    table_chunk = next(c for c in chunks if c.content_kind == ContentKind.TABLE)
    assert "Bait stations" in table_chunk.content
    assert table_chunk.location.object_id == "table-001"
    image_chunk = next(c for c in chunks if c.content_kind == ContentKind.IMAGE)
    assert image_chunk.location.object_id == "image-001"
    assert Path(image_chunk.metadata["extracted_path"]).exists()
    assert all(c.source_sha256 == sha256_file(path) for c in chunks)


def test_xlsx_parser_preserves_sheet_cell_ranges_formulas_and_blanks(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Pest Activity Log"
    sheet.append(["Date", "Count", "Assessment"])
    sheet.append(["2026-01-01", 3, "=B2>0"])
    sheet.append(["2026-02-01", None, "Clear"])
    workbook.save(path)

    chunks = parse_xlsx(_source(path, SourceType.XLSX, track=3), max_rows=20)

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.location.sheet == "Pest Activity Log"
    assert chunk.location.cell_range == "A1:C3"
    assert "<BLANK>" in chunk.content
    assert "=B2>0" in chunk.content
    assert chunk.metadata["formula_cells"] == ["C2"]


def test_xlsx_parser_flags_embedded_re_number_mismatch(tmp_path: Path) -> None:
    path = tmp_path / "mismatched.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["RE Number", "RE-VIC-2021-8004"])
    workbook.save(path)

    chunks = parse_xlsx(_source(path, SourceType.XLSX, track=3))

    assert chunks[0].re_number == "RE-WA-2021-0041"
    assert "embedded_re_number_mismatch" in chunks[0].flags
    assert chunks[0].metadata["embedded_re_numbers"] == ["RE-VIC-2021-8004"]


def test_pdf_fallback_preserves_page_numbers_and_records_mineru_unavailable(
    tmp_path: Path,
) -> None:
    path = tmp_path / "policy.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "Part 1 Registration requirements")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "Part 2 Record keeping requirements")
    pdf.save(path)
    pdf.close()
    source = SourceRecord(
        source_id="policy-rules-2021",
        path=path,
        source_type=SourceType.PDF,
        sha256=sha256_file(path),
    )

    chunks = parse_pdf(source, tmp_path / "mineru", mineru_executable=None)

    assert [chunk.location.page for chunk in chunks] == [1, 2]
    assert "Registration" in chunks[0].content
    assert all("mineru_unavailable" in chunk.flags for chunk in chunks)
