from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from freca.models import (
    ContentKind,
    EvidenceChunk,
    SourceLocation,
    SourceRecord,
)
from freca.parsing.chunking import normalize_text, stable_chunk_id
from freca.parsing.images import VisionDescriber, safe_description


def _base_chunk(
    source: SourceRecord,
    *,
    locator: str,
    content: str,
    kind: ContentKind,
    location: SourceLocation,
    derived_from: str | None = None,
    flags: list[str] | None = None,
    metadata: dict | None = None,
) -> EvidenceChunk:
    return EvidenceChunk(
        chunk_id=stable_chunk_id(source, locator),
        case_id=source.case_id,
        re_number=source.re_number,
        track=source.track,
        source_id=source.source_id,
        source_file=source.path.name,
        source_type=source.source_type,
        location=location,
        content=content,
        content_kind=kind,
        derived_from=derived_from,
        parser_name="python-docx+ooxml",
        parser_version="1",
        source_sha256=source.sha256,
        flags=list(source.flags) + list(flags or []),
        metadata=metadata or {},
    )


def parse_docx(
    source: SourceRecord,
    image_dir: Path,
    *,
    vision_describer: VisionDescriber | None = None,
) -> list[EvidenceChunk]:
    document = Document(source.path)
    chunks: list[EvidenceChunk] = []
    for index, paragraph in enumerate(document.paragraphs):
        content = normalize_text(paragraph.text)
        if not content:
            continue
        is_heading = paragraph.style is not None and paragraph.style.name.startswith("Heading")
        kind = ContentKind.HEADING if is_heading else ContentKind.PARAGRAPH
        chunks.append(
            _base_chunk(
                source,
                locator=f"paragraph-{index:04d}",
                content=content,
                kind=kind,
                location=SourceLocation(
                    section=paragraph.style.name if is_heading else None,
                    paragraph_index=index,
                ),
            )
        )

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(normalize_text(cell.text) or "<BLANK>" for cell in row.cells))
        content = "\n".join(rows)
        chunks.append(
            _base_chunk(
                source,
                locator=f"table-{table_index:03d}",
                content=content,
                kind=ContentKind.TABLE,
                location=SourceLocation(object_id=f"table-{table_index:03d}"),
                metadata={"row_count": len(table.rows), "column_count": len(table.columns)},
            )
        )

    image_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source.path) as archive:
        media = sorted(
            name for name in archive.namelist() if name.casefold().startswith("word/media/")
        )
        for image_index, member in enumerate(media, start=1):
            suffix = Path(member).suffix.casefold() or ".bin"
            extracted = image_dir / f"{source.source_id}_image-{image_index:03d}{suffix}"
            extracted.write_bytes(archive.read(member))
            object_id = f"image-{image_index:03d}"
            image_chunk = _base_chunk(
                source,
                locator=object_id,
                content=f"[Embedded image retained: {extracted.name}]",
                kind=ContentKind.IMAGE,
                location=SourceLocation(object_id=object_id),
                flags=["vision_description_pending"],
                metadata={"extracted_path": str(extracted.resolve()), "ooxml_member": member},
            )
            chunks.append(image_chunk)
            description = safe_description(
                vision_describer,
                extracted,
                context=f"{source.source_id} {object_id}",
            )
            if description:
                chunks.append(
                    _base_chunk(
                        source,
                        locator=f"{object_id}-description",
                        content=description,
                        kind=ContentKind.IMAGE_DESCRIPTION,
                        location=SourceLocation(object_id=object_id),
                        derived_from=image_chunk.chunk_id,
                        flags=["model_generated_neutral_description"],
                        metadata={"extracted_path": str(extracted.resolve())},
                    )
                )
    return chunks
