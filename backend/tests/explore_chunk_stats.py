"""
探索性统计：农场文档段落长度分布 + 不同 chunk_size 切分模拟
目的：用真实数据辅助选定 chunk_size（当前默认 400，待确认）
纯统计，不依赖 Docling，用降级解析（python-docx/openpyxl）快速跑。
"""
import os, re, random
from collections import defaultdict
from docx import Document
import openpyxl
import numpy as np

BASE = r"D:\桌面\农场任务二\Task2"
CASES_DIR = os.path.join(BASE, "SFRE_cases", "SFRE_cases")


def parse_docx(path):
    d = Document(path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def parse_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    out = []
    for ws in wb.worksheets:
        for r in ws.iter_rows(values_only=True):
            cells = [str(v) for v in r if v is not None and str(v).strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def parse_file(path):
    if path.lower().endswith(".docx"):
        return parse_docx(path)
    if path.lower().endswith(".xlsx"):
        return parse_xlsx(path)
    return ""


def split_sents(text):
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+|\n+", text) if s.strip()]


def pct(a, p):
    return int(np.percentile(a, p))


# 抽样：3 个异常 + 随机 5 个
all_cases = sorted(c for c in os.listdir(CASES_DIR) if c.startswith("RE-"))
must = ["RE-QLD-2022-0077", "RE-SA-2021-0066", "RE-WA-2021-0077"]
rest = [c for c in all_cases if c not in must]
random.seed(42)
sample = must + random.sample(rest, 5)

# 收集每个文件的句子
files_sents = []  # (cid, track, file, [sents])
track_sent = defaultdict(list)
for cid in sample:
    cdir = os.path.join(CASES_DIR, cid)
    for f in sorted(os.listdir(cdir)):
        fp = os.path.join(cdir, f)
        if not os.path.isfile(fp):
            continue
        txt = parse_file(fp)
        if not txt.strip():
            continue
        m = re.match(r"(\d+)_", f)
        track = m.group(1) if m else "?"
        sents = split_sents(txt)
        files_sents.append((cid, track, f, sents))
        for s in sents:
            track_sent[track].append(len(s))

sent_lens = np.array([len(s) for (_, _, _, ss) in files_sents for s in ss])

print("=" * 64)
print(f"抽样 {len(sample)} 个 case（含 3 异常），共 {len(files_sents)} 个文件")
print(f"句子总数: {len(sent_lens)}")
print("=" * 64)
print("【句子字符长度分布】")
print(f"  min={sent_lens.min()} P25={pct(sent_lens,25)} P50={pct(sent_lens,50)} "
      f"P75={pct(sent_lens,75)} P90={pct(sent_lens,90)} P95={pct(sent_lens,95)} "
      f"P99={pct(sent_lens,99)} max={sent_lens.max()} mean={int(sent_lens.mean())}")
print(f"  token近似(字符/4): P50={pct(sent_lens,50)//4} P90={pct(sent_lens,90)//4} "
      f"P99={pct(sent_lens,99)//4} max={sent_lens.max()//4}")
for thr in [400, 600, 800, 1200, 2000]:
    n = int((sent_lens > thr).sum())
    print(f"  句子 >{thr}字符(~{thr//4}tok): {n} ({n/len(sent_lens)*100:.1f}%)")

print()
print("【按 track 句子字符长度】P50 / P90 / max / 句子数")
for t in sorted(track_sent):
    a = np.array(track_sent[t])
    print(f"  Track{t}: P50={pct(a,50)} P90={pct(a,90)} max={a.max()} n={len(a)}")


def simulate(size):
    cl = []
    for (_, _, _, ss) in files_sents:
        cur = ""
        for s in ss:
            if len(cur) + len(s) <= size:
                cur = (cur + " " + s).strip()
            else:
                if cur:
                    cl.append(len(cur))
                cur = s
        if cur:
            cl.append(len(cur))
    return np.array(cl)


print()
print("【不同 chunk_size 切分模拟】(按句子整块累加, 不切断句子)")
print(f"  {'size字符':<10}{'~token':<8}{'chunk数':<8}{'P50':<7}{'P90':<7}{'max':<7}{'mean':<7}{'超长chunk>2*size':<14}")
for size in [400, 800, 1200, 1600, 2000, 2400]:
    ca = simulate(size)
    over = int((ca > 2 * size).sum())
    print(f"  {size:<10}{size//4:<8}{len(ca):<8}{pct(ca,50):<7}{pct(ca,90):<7}"
          f"{ca.max():<7}{int(ca.mean()):<7}{over}")

print()
print("【关键发现】")
print(f"  - 句子 P90={pct(sent_lens,90)}字符(~{pct(sent_lens,90)//4}tok), "
      f"P99={pct(sent_lens,99)}字符(~{pct(sent_lens,99)//4}tok), max={sent_lens.max()}字符(~{sent_lens.max()//4}tok)")
print(f"  - 当前 chunk_size=400(字符) ≈ 100 token，远小于常见 400 token，切得过细")
s400 = simulate(400)
print(f"  - size=400 切出 {len(s400)} 个 chunk, 平均 {int(s400.mean())} 字符(~{int(s400.mean())//4} tok/chunk)")
