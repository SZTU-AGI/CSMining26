"""降级解析器：Docling 不可用（未安装/下载失败）时启用轻量库。

仅用于兜底，不依赖任何 ML 模型。
"""
from docx import Document
import openpyxl


def fallback_parse_docx(path: str) -> str:
    d = Document(path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def fallback_parse_xlsx(path: str) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    out = []
    for ws in wb.worksheets:
        for r in ws.iter_rows(values_only=True):
            cells = [str(v) for v in r if v is not None and str(v).strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def fallback_parse_pdf(path: str) -> str:
    import pdfplumber  # 惰性导入: 农场证据均 docx/xlsx, 仅解析 PDF 时才需要
    out = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                out.append(txt)
    return "\n".join(out)


def fallback_parse(path: str) -> str:
    p = path.lower()
    if p.endswith(".docx"):
        return fallback_parse_docx(path)
    if p.endswith(".xlsx"):
        return fallback_parse_xlsx(path)
    if p.endswith(".pdf"):
        return fallback_parse_pdf(path)
    return ""
