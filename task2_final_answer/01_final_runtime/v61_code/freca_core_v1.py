from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import sys
import time
from collections import Counter
from pathlib import Path
from typing import Any

import pymupdf
import requests
from docx import Document
from openpyxl import load_workbook


# ============================================================
# Paths
# ============================================================

TASK_ROOT = Path(
    os.environ.get(
        "FRECA_TASK_ROOT",
        "/home/MeggieYu/freca/Task2",
    )
)

PROJECT_ROOT = Path(
    os.environ.get(
        "FRECA_PROJECT_ROOT",
        "/home/MeggieYu/freca/core_v1",
    )
)

RULES_PATH = (
    TASK_ROOT
    / "1-Export Control (Plants and Plant Products)Rules 2021.pdf"
)

CP_PATH = (
    TASK_ROOT
    / "checkingpoints_all_elements_onesheet.xlsx"
)

CONTRACT_DIR = PROJECT_ROOT / "contracts"
RESULT_DIR = PROJECT_ROOT / "results"

API_PROVIDER = os.environ.get("FRECA_API_PROVIDER", "deepseek").lower()
API_BASE_URL = (
    os.environ.get("ZHIPU_BASE_URL")
    or os.environ.get("ZAI_BASE_URL")
    or "https://open.bigmodel.cn/api/paas/v4"
    if API_PROVIDER in {"zhipu", "zai", "bigmodel"}
    else os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
)
DEEPSEEK_BASE_URL = API_BASE_URL

CONTRACT_MODEL = os.environ.get(
    "FRECA_CONTRACT_MODEL",
    "glm-4.5-air" if API_PROVIDER in {"zhipu", "zai", "bigmodel"}
    else "deepseek-v4-pro",
)

ALIGNMENT_MODEL = os.environ.get(
    "FRECA_ALIGNMENT_MODEL",
    "glm-4.5-air" if API_PROVIDER in {"zhipu", "zai", "bigmodel"}
    else "deepseek-v4-flash",
)
EVIDENCE_MODEL = ALIGNMENT_MODEL


# ============================================================
# Utilities
# ============================================================

def normalize_text(value: Any) -> str:
    if value is None:
        return ""

    text = str(value)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def canonical_cp_id(value: str) -> str:
    return re.sub(
        r"[^A-Za-z0-9]+",
        "",
        str(value),
    ).upper()


def quote_match_text(value: Any) -> str:
    """
    Normalize only whitespace for source-quote validation.

    PDF line wrapping such as:
        "prescribed\\nplants"
    should match:
        "prescribed plants"

    No words, punctuation, or characters are otherwise changed.
    """
    return re.sub(
        r"\\s+",
        " ",
        str(value),
    ).strip()


def tokenize(text: str) -> list[str]:
    return re.findall(
        r"[a-z0-9]+|[\u4e00-\u9fff]",
        text.lower(),
    )


def save_json(data: Any, path: Path):
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    path.write_text(
        json.dumps(
            data,
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def load_json(path: Path) -> Any:
    return json.loads(
        path.read_text(
            encoding="utf-8"
        )
    )


# ============================================================
# Lightweight BM25
# ============================================================

def bm25_rank(
    query: str,
    docs: list[dict],
    top_k: int,
) -> list[dict]:

    if not docs:
        return []

    query_tokens = tokenize(query)

    tokenized_docs = [
        tokenize(doc["text"])
        for doc in docs
    ]

    n_docs = len(docs)

    avg_dl = (
        sum(len(tokens) for tokens in tokenized_docs)
        / max(n_docs, 1)
    )

    document_frequency = Counter()

    for tokens in tokenized_docs:
        for token in set(tokens):
            document_frequency[token] += 1

    k1 = 1.5
    b = 0.75

    ranked = []

    for doc, tokens in zip(
        docs,
        tokenized_docs,
    ):
        tf = Counter(tokens)
        dl = len(tokens)

        score = 0.0

        for term in query_tokens:

            if term not in tf:
                continue

            df = document_frequency[term]

            idf = math.log(
                1
                + (
                    n_docs
                    - df
                    + 0.5
                )
                / (
                    df
                    + 0.5
                )
            )

            freq = tf[term]

            denominator = (
                freq
                + k1
                * (
                    1
                    - b
                    + b
                    * dl
                    / max(avg_dl, 1)
                )
            )

            score += (
                idf
                * freq
                * (k1 + 1)
                / denominator
            )

        item = dict(doc)
        item["score"] = score

        ranked.append(item)

    ranked.sort(
        key=lambda x: (
            -x["score"],
            x["id"],
        )
    )

    return ranked[:top_k]


# ============================================================
# Official CP workbook
# ============================================================

def load_official_cps() -> list[dict]:

    wb = load_workbook(
        CP_PATH,
        data_only=True,
    )

    result = []

    for ws in wb.worksheets:

        current_element = ""
        current_subelement = ""

        for col in range(
            1,
            ws.max_column + 1,
        ):

            row1 = normalize_text(
                ws.cell(1, col).value
            )

            row2 = normalize_text(
                ws.cell(2, col).value
            )

            criterion = normalize_text(
                ws.cell(3, col).value
            )

            cp_id = normalize_text(
                ws.cell(4, col).value
            )

            if row1:
                current_element = row1

            if row2:
                current_subelement = row2

            if not cp_id:
                continue

            result.append(
                {
                    "cp_id": cp_id,
                    "canonical_cp_id":
                        canonical_cp_id(cp_id),
                    "element":
                        current_element,
                    "subelement":
                        current_subelement,
                    "criterion":
                        criterion,
                    "sheet":
                        ws.title,
                    "column":
                        col,
                }
            )

    return result


def get_cp(cp_id: str) -> dict:

    target = canonical_cp_id(cp_id)

    cps = load_official_cps()

    matches = [
        cp
        for cp in cps
        if cp["canonical_cp_id"]
        == target
    ]

    if len(matches) != 1:
        available = ", ".join(
            cp["cp_id"]
            for cp in cps
        )

        raise ValueError(
            f"Cannot uniquely find {cp_id}.\n"
            f"Available CPs:\n{available}"
        )

    return matches[0]


# ============================================================
# Rules PDF
# ============================================================

def extract_rule_chunks(
    max_chars: int = 1500,
) -> list[dict]:

    pdf = pymupdf.open(
        RULES_PATH
    )

    chunks = []

    for page_index in range(
        len(pdf)
    ):

        text = pdf[
            page_index
        ].get_text("text")

        lines = [
            normalize_text(line)
            for line
            in text.splitlines()
            if normalize_text(line)
        ]

        buffer = []
        count = 0
        chunk_number = 0

        def flush():
            nonlocal buffer
            nonlocal count
            nonlocal chunk_number

            if not buffer:
                return

            chunk_number += 1

            chunks.append(
                {
                    "id":
                        f"RULE-P{page_index + 1}"
                        f"-C{chunk_number}",
                    "page":
                        page_index + 1,
                    "text":
                        "\n".join(buffer),
                }
            )

            buffer = []
            count = 0

        for line in lines:

            if (
                buffer
                and count + len(line)
                > max_chars
            ):
                flush()

            buffer.append(line)
            count += len(line) + 1

        flush()

    return chunks


# ============================================================
# Evidence files
# ============================================================

def parse_docx(
    path: Path,
) -> list[dict]:

    doc = Document(path)

    chunks = []

    paragraph_number = 0

    for paragraph in doc.paragraphs:

        text = normalize_text(
            paragraph.text
        )

        if not text:
            continue

        paragraph_number += 1

        chunks.append(
            {
                "id":
                    f"{path.name}:P"
                    f"{paragraph_number}",
                "file":
                    path.name,
                "kind":
                    "DOCX_PARAGRAPH",
                "text":
                    text,
            }
        )

    for table_number, table in enumerate(
        doc.tables,
        1,
    ):
        for row_number, row in enumerate(
            table.rows,
            1,
        ):

            values = [
                normalize_text(cell.text)
                for cell in row.cells
            ]

            if not any(values):
                continue

            chunks.append(
                {
                    "id":
                        f"{path.name}:"
                        f"T{table_number}:"
                        f"R{row_number}",
                    "file":
                        path.name,
                    "kind":
                        "DOCX_TABLE_ROW",
                    "text":
                        " | ".join(values),
                }
            )

    return chunks


def parse_xlsx(
    path: Path,
) -> list[dict]:

    wb = load_workbook(
        path,
        read_only=True,
        data_only=False,
    )

    chunks = []

    for ws in wb.worksheets:

        for row_number, row in enumerate(
            ws.iter_rows(
                values_only=True
            ),
            1,
        ):

            values = [
                normalize_text(value)
                for value in row
            ]

            if not any(values):
                continue

            chunks.append(
                {
                    "id":
                        f"{path.name}:"
                        f"{ws.title}:"
                        f"R{row_number}",
                    "file":
                        path.name,
                    "sheet":
                        ws.title,
                    "kind":
                        "XLSX_ROW",
                    "text":
                        " | ".join(values),
                }
            )

    return chunks


def find_case_dir(
    case_name: str,
) -> Path:

    exact = [
        path
        for path in TASK_ROOT.rglob(
            case_name
        )
        if path.is_dir()
        and path.name == case_name
    ]

    if not exact:
        raise FileNotFoundError(
            f"Cannot find case directory "
            f"{case_name} under "
            f"{TASK_ROOT}"
        )

    # Sometimes extracted ZIPs introduce
    # an extra SFRE_cases level.
    # Recursive discovery avoids hard-coding it.
    if len(exact) > 1:

        print(
            "[WARN] Multiple matching "
            "case directories found:"
        )

        for path in exact:
            print("   ", path)

        print(
            "[WARN] Using the shortest path."
        )

        exact.sort(
            key=lambda x: (
                len(x.parts),
                str(x),
            )
        )

    return exact[0]


def load_case_evidence(
    case_dir: Path,
) -> list[dict]:

    files = sorted(
        [
            path
            for path
            in case_dir.rglob("*")
            if path.is_file()
            and path.suffix.lower()
            in {".docx", ".xlsx"}
            and not path.name.startswith(
                "~$"
            )
        ],
        key=lambda x: str(x),
    )

    chunks = []

    for path in files:

        print(
            f"    parsing: {path.name}"
        )

        try:

            if (
                path.suffix.lower()
                == ".docx"
            ):
                chunks.extend(
                    parse_docx(path)
                )

            elif (
                path.suffix.lower()
                == ".xlsx"
            ):
                chunks.extend(
                    parse_xlsx(path)
                )

        except Exception as exc:

            print(
                f"[WARN] Cannot parse "
                f"{path}: {exc}",
                file=sys.stderr,
            )

    return chunks


# ============================================================
# DeepSeek API
# ============================================================

def extract_json(
    text: str,
) -> dict:

    if not text:
        raise ValueError(
            "Model returned empty content"
        )

    text = text.strip()

    text = re.sub(
        r"<think>.*?</think>",
        "",
        text,
        flags=re.S,
    ).strip()

    text = re.sub(
        r"^```(?:json)?\s*",
        "",
        text,
        flags=re.I,
    )

    text = re.sub(
        r"\s*```$",
        "",
        text,
    )

    try:
        return json.loads(text)

    except json.JSONDecodeError:

        start = text.find("{")
        end = text.rfind("}")

        if (
            start >= 0
            and end > start
        ):
            return json.loads(
                text[start:end + 1]
            )

        raise


def deepseek_json(
    *,
    model: str,
    system_prompt: str,
    user_prompt: str,
    thinking: bool,
    max_tokens: int,
) -> dict:

    is_zhipu = API_PROVIDER in {"zhipu", "zai", "bigmodel"}
    api_key = (
        os.environ.get("ZHIPU_API_KEY")
        or os.environ.get("ZAI_API_KEY")
        if is_zhipu
        else os.environ.get("DEEPSEEK_API_KEY")
    )

    if not api_key:
        raise RuntimeError(
            ("ZHIPU_API_KEY/ZAI_API_KEY" if is_zhipu else "DEEPSEEK_API_KEY")
            + " is not set."
        )

    endpoint = (
        API_BASE_URL.rstrip("/")
        + "/chat/completions"
    )

    headers = {
        "Authorization":
            f"Bearer {api_key}",
        "Content-Type":
            "application/json",
    }

    payload = {
        "model":
            model,
        "messages":
            [
                {
                    "role":
                        "system",
                    "content":
                        system_prompt,
                },
                {
                    "role":
                        "user",
                    "content":
                        user_prompt,
                },
            ],
        "stream":
            False,
        "max_tokens":
            max_tokens,
        "response_format":
            {
                "type":
                    "json_object"
            },
        "thinking":
            {
                "type":
                    (
                        "enabled"
                        if thinking
                        else "disabled"
                    )
            },
    }

    if thinking:
        payload[
            "reasoning_effort"
        ] = "high"

    # Pin decoding on every branch.  Contracts and alignments are persisted
    # and hashed, but an omitted decoding parameter still makes a clean-room
    # rebuild depend on the provider default.  The escape hatch is explicit
    # and must be captured by the V6 run manifest.
    if not (
        thinking
        and os.environ.get("FRECA_UNPIN_THINKING_TEMPERATURE") == "1"
    ):
        payload[
            "temperature"
        ] = 0.0

    last_error = None

    max_attempts = int(os.environ.get("FRECA_API_MAX_ATTEMPTS", "6"))
    for attempt in range(1, max_attempts + 1):

        try:

            print(
                f"    API call "
                f"{model} "
                f"(attempt {attempt}/{max_attempts})"
            )

            # A syntactically malformed JSON response is often reproduced
            # byte-for-byte when temperature=0 and the retry sends the exact
            # same prompt (provider caching can reinforce this).  Preserve
            # deterministic decoding, but make retry attempts a distinct
            # request and explicitly ask for syntax correction.
            request_payload = dict(payload)
            request_payload["messages"] = [
                dict(message)
                for message in payload["messages"]
            ]
            if attempt > 1:
                request_payload["messages"][1]["content"] = (
                    user_prompt
                    + "\n\nSTRICT_JSON_RETRY_ATTEMPT_"
                    + str(attempt)
                    + ": The previous provider response was not valid JSON. "
                      "Return the same semantic answer as exactly one valid JSON "
                      "object. Check every comma, quote, bracket, and brace. "
                      "Do not use Markdown fences or explanatory text."
                )

            response = requests.post(
                endpoint,
                headers=headers,
                json=request_payload,
                timeout=300,
            )

            try:
                response.raise_for_status()
            except requests.HTTPError as http_exc:
                # Preserve provider diagnostics without headers or secrets.
                status = response.status_code
                body = re.sub(r"(?i)(bearer\\s+)[^\\s]+", r"\\1<redacted>", response.text)
                body = body[:1200]
                http_exc.args = (f"HTTP {status}: {body}",)
                raise

            data = response.json()

            message = (
                data["choices"][0]
                ["message"]
            )

            content = (
                message.get("content")
                or ""
            )

            if not content.strip():

                reasoning_content = (
                    message.get(
                        "reasoning_content"
                    )
                    or ""
                )

                finish_reason = (
                    data["choices"][0]
                    .get("finish_reason")
                )

                usage = data.get(
                    "usage",
                    {}
                )

                print(
                    "[WARN] Empty content details:",
                    {
                        "finish_reason":
                            finish_reason,
                        "reasoning_chars":
                            len(
                                reasoning_content
                            ),
                        "usage":
                            usage,
                    },
                    file=sys.stderr,
                )

                raise ValueError(
                    "Empty model content"
                )

            finish_reason = (
                data["choices"][0]
                .get("finish_reason")
            )

            if finish_reason == "length":
                raise ValueError(
                    "Model output was truncated "
                    "(finish_reason=length)."
                )

            try:
                result = extract_json(
                    content
                )

            except Exception as parse_exc:

                debug_dir = (
                    PROJECT_ROOT
                    / "api_debug"
                )

                debug_dir.mkdir(
                    parents=True,
                    exist_ok=True,
                )

                digest = hashlib.sha256(
                    content.encode(
                        "utf-8",
                        errors="replace",
                    )
                ).hexdigest()[:12]

                safe_model = re.sub(
                    r"[^A-Za-z0-9_.-]+",
                    "_",
                    model,
                )

                debug_path = (
                    debug_dir
                    / (
                        f"{safe_model}_"
                        f"json_error_"
                        f"{digest}.txt"
                    )
                )

                debug_path.write_text(
                    content,
                    encoding="utf-8",
                )

                print(
                    f"[WARN] Raw malformed JSON "
                    f"saved to: {debug_path}",
                    file=sys.stderr,
                )

                raise ValueError(
                    f"JSON parse failed: "
                    f"{parse_exc}"
                ) from parse_exc

            usage = data.get(
                "usage",
                {}
            )

            if usage:
                print(
                    "    tokens:",
                    usage,
                )

            return result

        except Exception as exc:

            last_error = exc

            print(
                f"[WARN] API attempt "
                f"{attempt} failed: "
                f"{exc}",
                file=sys.stderr,
            )

            status = getattr(getattr(exc, "response", None), "status_code", None)
            retryable_status = status is None or status == 429 or status >= 500
            if attempt < max_attempts and retryable_status:
                delay = min(60, 2 ** (attempt - 1))
                retry_after = None
                response_obj = getattr(exc, "response", None)
                if response_obj is not None:
                    retry_after = response_obj.headers.get("Retry-After")
                if retry_after:
                    try:
                        delay = min(300, max(delay, float(retry_after)))
                    except ValueError:
                        pass
                time.sleep(delay)

    raise RuntimeError(
        f"{API_PROVIDER} API failed: "
        f"{last_error}"
    )


# ============================================================
# Contract compiler
# ============================================================

CONTRACT_SYSTEM = r"""
You are a legal rule compiler.

You are NOT a compliance classifier.

You receive:
1. exactly one official FRECA checking point (CP criterion);
2. retrieved excerpts from the official Export Control Rules.

Your task is to compile the SMALLEST legally grounded executable
contract needed to evaluate that specific CP criterion.

CRITICAL PRINCIPLE:

RELEVANCE IS NOT THE SAME AS A SCORING REQUIREMENT.

The retrieved Rules excerpts may contain many legally related duties.
Do NOT put every related duty into the CP satisfaction logic.

The official CP criterion defines what is being scored.
Rules excerpts provide legal grounding and interpretation for that
criterion.

Before including any proposition in satisfaction, apply this test:

    If this proposition were false, would that by itself mean that the
    supplied CP criterion is not satisfied?

If the answer is no, unclear, or merely "this rule is related to the
same topic", do NOT make it a mandatory satisfaction atom.

Classify the role of each included proposition conceptually as one of:

- DIRECT_REQUIREMENT:
  directly defines or operationalises wording in the supplied CP criterion.

- APPLICABILITY:
  determines whether the whole CP applies.

- NON_APPLICABILITY:
  positively establishes that the whole CP does not apply.

Do NOT create mandatory atoms from provisions that are merely:

- SUPPORTING_CONTEXT
- RELATED_NOT_SCORED
- a separate operational obligation not required by this CP criterion
- another compliance requirement from the same regulatory section

MINIMALITY RULE:

Prefer the smallest non-redundant set of propositions that faithfully
represents the CP criterion.

Do not split one holistic criterion into multiple mandatory atoms unless
the official wording or Rules structure genuinely requires all of them.

Do not convert examples, supporting controls, or related operational
requirements into additional mandatory conjuncts.

Every ATOM must:

1. represent a testable factual proposition;
2. directly correspond to wording in the supplied CP criterion;
3. contain a short exact "cp_quote" copied verbatim from the CP criterion;
4. contain at least one exact legal anchor copied from the supplied
   CP text or Rules excerpt.

Use only supplied material.
Never use outside knowledge.
Never decide any specific case.
Never output 1, 0, N/A, PASS, FAIL, COMPLIANT or NON-COMPLIANT.

Missing evidence is NOT non-applicability.

Non-applicability must be positively established.

Return valid JSON in this structure:

{
  "cp_id": "CP12",

  "atoms": [
    {
      "atom_id": "A1",
      "proposition": "A testable factual proposition",
      "cp_quote": "exact phrase from the CP criterion",
      "anchors": [
        {
          "source": "RULES",
          "chunk_id": "RULE-P1-C1",
          "quote": "exact quote copied from the Rules chunk"
        }
      ]
    }
  ],

  "applicability": {
    "op": "CONST",
    "value": true
  },

  "satisfaction": {
    "op": "ATOM",
    "atom_id": "A1"
  },

  "non_applicability": {
    "op": "CONST",
    "value": false
  },

  "related_not_scored": [
    {
      "chunk_id": "RULE-P1-C2",
      "reason": "Related legal duty but not independently required by this CP criterion."
    }
  ],

  "notes": []
}

Allowed expressions:

{"op":"ATOM","atom_id":"A1"}

{"op":"ALL","children":[...]}

{"op":"ANY","children":[...]}

{"op":"NOT","children":[one_expression]}

{"op":"CONST","value":true}

{"op":"CONST","value":false}

Use ALL only when the CP criterion truly requires every child.

Use ANY only when the official material supplies genuine alternative
ways of satisfying the criterion.

applicability:
positive conditions required before this CP applies.
If there is no special whole-CP applicability condition supported by
the supplied material, use CONST true.

satisfaction:
only the direct requirements necessary to satisfy THIS CP criterion.

non_applicability:
only positive conditions establishing that the whole CP does not apply.
If none are supported, use CONST false.

related_not_scored:
record retrieved provisions that are relevant to the topic but should
not become independent scoring conditions.

Return JSON only.
"""


def make_contract_prompt(
    cp: dict,
    rule_chunks: list[dict],
) -> str:

    rule_text = "\n\n".join(
        (
            f"[{chunk['id']}] "
            f"PAGE={chunk['page']}\n"
            f"{chunk['text']}"
        )
        for chunk in rule_chunks
    )

    return f"""
OFFICIAL CHECKING POINT

CP_ID:
{cp["cp_id"]}

ELEMENT:
{cp["element"]}

SUBELEMENT:
{cp["subelement"]}

CRITERION:
{cp["criterion"]}


RETRIEVED OFFICIAL RULES CHUNKS

{rule_text}


Compile this checking point.

The response must be valid JSON.
"""


def validate_expression(
    expression: dict,
    valid_atom_ids: set[str],
):

    if not isinstance(
        expression,
        dict,
    ):
        raise ValueError(
            "Expression must be object"
        )

    op = expression.get("op")

    if op == "CONST":

        if not isinstance(
            expression.get("value"),
            bool,
        ):
            raise ValueError(
                "CONST requires bool value"
            )

        return

    if op == "ATOM":

        atom_id = expression.get(
            "atom_id"
        )

        if atom_id not in valid_atom_ids:
            raise ValueError(
                f"Unknown atom: {atom_id}"
            )

        return

    if op not in {
        "ALL",
        "ANY",
        "NOT",
    }:
        raise ValueError(
            f"Unknown op: {op}"
        )

    children = expression.get(
        "children"
    )

    if not isinstance(
        children,
        list,
    ):
        raise ValueError(
            f"{op} requires children"
        )

    if (
        op == "NOT"
        and len(children) != 1
    ):
        raise ValueError(
            "NOT requires exactly "
            "one child"
        )

    if (
        op in {"ALL", "ANY"}
        and len(children) == 0
    ):
        raise ValueError(
            f"{op} cannot be empty"
        )

    for child in children:
        validate_expression(
            child,
            valid_atom_ids,
        )


def validate_contract(
    contract: dict,
    cp: dict,
    rule_chunks: list[dict],
):

    if (
        canonical_cp_id(
            contract.get(
                "cp_id",
                "",
            )
        )
        != cp[
            "canonical_cp_id"
        ]
    ):
        raise ValueError(
            "Contract CP ID mismatch"
        )

    atoms = contract.get(
        "atoms"
    )

    if (
        not isinstance(atoms, list)
        or not atoms
    ):
        raise ValueError(
            "Contract has no atoms"
        )

    rule_map = {
        chunk["id"]:
            chunk["text"]
        for chunk in rule_chunks
    }

    atom_ids = set()

    for atom in atoms:

        atom_id = normalize_text(
            atom.get("atom_id")
        )

        proposition = normalize_text(
            atom.get("proposition")
        )

        cp_quote = normalize_text(
            atom.get("cp_quote")
        )

        if not cp_quote:
            raise ValueError(
                f"{atom_id}: missing cp_quote"
            )

        if (
            quote_match_text(cp_quote)
            not in quote_match_text(
                cp["criterion"]
            )
        ):
            raise ValueError(
                f"{atom_id}: cp_quote is not "
                f"verbatim from criterion:\n"
                f"{cp_quote}"
            )

        if not atom_id:
            raise ValueError(
                "Atom without ID"
            )

        if atom_id in atom_ids:
            raise ValueError(
                f"Duplicate atom ID "
                f"{atom_id}"
            )

        atom_ids.add(atom_id)

        if not proposition:
            raise ValueError(
                f"{atom_id} has no "
                f"proposition"
            )

        anchors = atom.get(
            "anchors"
        )

        if (
            not isinstance(
                anchors,
                list,
            )
            or not anchors
        ):
            raise ValueError(
                f"{atom_id} has no "
                f"anchors"
            )

        for anchor in anchors:

            source = anchor.get(
                "source"
            )

            quote = normalize_text(
                anchor.get("quote")
            )

            if not quote:
                raise ValueError(
                    f"{atom_id}: "
                    f"empty quote"
                )

            if source == "CP":

                if (
                    quote_match_text(quote)
                    not in quote_match_text(
                        cp["criterion"]
                    )
                ):
                    raise ValueError(
                        f"{atom_id}: "
                        f"CP quote is not "
                        f"verbatim:\n{quote}"
                    )

            elif source == "RULES":

                chunk_id = anchor.get(
                    "chunk_id"
                )

                if (
                    chunk_id
                    not in rule_map
                ):
                    raise ValueError(
                        f"{atom_id}: "
                        f"unknown Rules "
                        f"chunk {chunk_id}"
                    )

                if (
                    quote_match_text(quote)
                    not in quote_match_text(
                        rule_map[chunk_id]
                    )
                ):
                    raise ValueError(
                        f"{atom_id}: "
                        f"Rules quote is "
                        f"not verbatim:\n"
                        f"{quote}"
                    )

            else:
                raise ValueError(
                    f"{atom_id}: "
                    f"invalid source "
                    f"{source}"
                )

    for root in (
        "applicability",
        "satisfaction",
        "non_applicability",
    ):

        if root not in contract:
            raise ValueError(
                f"Missing root: {root}"
            )

        validate_expression(
            contract[root],
            atom_ids,
        )


def compile_cp(
    cp_id: str,
    policy_top_k: int,
) -> Path:

    print("\n" + "=" * 72)
    print("FRECA CORE V1 — COMPILE")
    print("=" * 72)

    cp = get_cp(cp_id)

    print(
        f"\nCP: {cp['cp_id']}"
    )
    print(
        f"Criterion:\n"
        f"{cp['criterion']}"
    )

    print(
        "\n[1/3] Parsing Rules PDF..."
    )

    all_rules = extract_rule_chunks()

    print(
        f"Rules chunks: "
        f"{len(all_rules)}"
    )

    query = " ".join(
        [
            cp["element"],
            cp["subelement"],
            cp["criterion"],
        ]
    )

    print(
        "\n[2/3] BM25 Rules retrieval..."
    )

    retrieved = bm25_rank(
        query,
        all_rules,
        policy_top_k,
    )

    for chunk in retrieved[:10]:
        print(
            f"  {chunk['id']}: "
            f"{chunk['score']:.3f}"
        )

    print(
        "\n[3/3] Compiling "
        f"with {CONTRACT_MODEL}..."
    )

    contract = None
    validation_error = None

    for semantic_attempt in range(
        1,
        3,
    ):

        prompt = make_contract_prompt(
            cp,
            retrieved,
        )

        if validation_error:

            prompt += (
                "\n\nPREVIOUS OUTPUT "
                "FAILED VALIDATION.\n"
                "Validation error:\n"
                f"{validation_error}\n\n"
                "Regenerate the full JSON "
                "contract and correct this "
                "problem."
            )

        candidate = deepseek_json(
            model=CONTRACT_MODEL,
            system_prompt=
                CONTRACT_SYSTEM,
            user_prompt=
                prompt,
            thinking=True,
            max_tokens=6000,
        )

        try:

            validate_contract(
                candidate,
                cp,
                retrieved,
            )

            contract = candidate
            break

        except Exception as exc:

            validation_error = str(exc)

            print(
                "[WARN] Contract "
                "validation failed:",
                validation_error,
                file=sys.stderr,
            )

    if contract is None:
        raise RuntimeError(
            "Could not produce a "
            "valid contract."
        )

    output_path = (
        CONTRACT_DIR
        / f"{cp['cp_id']}.json"
    )

    save_json(
        {
            "schema":
                "freca-core-contract-v1",
            "cp":
                cp,
            "model":
                CONTRACT_MODEL,
            "retrieved_rules":
                retrieved,
            "contract":
                contract,
        },
        output_path,
    )

    print("\nCONTRACT COMPILED")
    print("-" * 72)

    for atom in contract["atoms"]:

        print(
            f"{atom['atom_id']}: "
            f"{atom['proposition']}"
        )

        for anchor in atom[
            "anchors"
        ]:
            print(
                "   source:",
                anchor.get("source"),
                anchor.get(
                    "chunk_id",
                    "",
                ),
            )
            print(
                "   quote:",
                anchor.get("quote"),
            )

    print(
        "\nSaved:"
        f"\n{output_path}"
    )

    return output_path


# ============================================================
# Evidence alignment
# ============================================================

ALIGNMENT_SYSTEM = r"""
You are an evidence-to-proposition alignment component.

You are NOT a compliance classifier.

For every supplied proposition, inspect only the supplied evidence
chunks from the current case.

Return evidence that SUPPORTS the proposition and evidence that
ATTACKS the proposition.

SUPPORT means that the source explicitly provides factual support
for the proposition.

ATTACK means that the source explicitly provides factual evidence
that contradicts the proposition.

If evidence is missing, vague, merely procedural, about another
entity, another establishment, another commodity, another time
period, or otherwise insufficient, do not manufacture support or
attack.

Missing evidence is UNKNOWN, not ATTACK.

Do not output a final compliance label.

Do not infer PASS, FAIL, 1, 0 or N/A.

Words written in a source such as compliant, non-compliant, pass,
fail, audit result, CP number, or N/A are source language, not your
final judgment.

Every cited quote must be copied verbatim from exactly one supplied
evidence chunk.

An atom can have both support and attack evidence.

Return valid JSON in this form:

{
  "alignments": [
    {
      "atom_id": "A1",
      "support": [
        {
          "chunk_id": "filename:P1",
          "quote": "exact quote"
        }
      ],
      "attack": [],
      "reason": "brief factual explanation"
    }
  ]
}

Return one alignment entry for every atom.

Return JSON only.
"""


def make_alignment_prompt(
    contract: dict,
    evidence: list[dict],
) -> str:

    propositions = [
        {
            "atom_id":
                atom["atom_id"],
            "proposition":
                atom["proposition"],
        }
        for atom
        in contract["atoms"]
    ]

    evidence_text = "\n\n".join(
        (
            f"[{chunk['id']}]\n"
            f"FILE={chunk['file']}\n"
            f"{chunk['text']}"
        )
        for chunk in evidence
    )

    return f"""
PROPOSITIONS

{json.dumps(
    propositions,
    ensure_ascii=False,
    indent=2
)}


CURRENT CASE EVIDENCE

{evidence_text}


Align every proposition with the case evidence.

The response must be valid JSON.
"""


def validate_alignments(
    raw: dict,
    atoms: list[dict],
    evidence: list[dict],
) -> dict[str, dict]:

    atom_ids = {
        atom["atom_id"]
        for atom in atoms
    }

    evidence_map = {
        chunk["id"]:
            chunk["text"]
        for chunk in evidence
    }

    result = {}

    for item in raw.get(
        "alignments",
        [],
    ):

        atom_id = item.get(
            "atom_id"
        )

        if atom_id not in atom_ids:
            continue

        support = []
        attack = []

        for source_field, target in (
            ("support", support),
            ("attack", attack),
        ):

            items = item.get(
                source_field,
                [],
            )

            if not isinstance(
                items,
                list,
            ):
                continue

            for evidence_item in items:

                chunk_id = (
                    evidence_item.get(
                        "chunk_id"
                    )
                )

                quote = normalize_text(
                    evidence_item.get(
                        "quote"
                    )
                )

                if (
                    chunk_id
                    in evidence_map
                    and quote
                    and quote_match_text(quote)
                    in quote_match_text(
                        evidence_map[chunk_id]
                    )
                ):
                    target.append(
                        {
                            "chunk_id":
                                chunk_id,
                            "quote":
                                quote,
                        }
                    )

        result[atom_id] = {
            "atom_id":
                atom_id,
            "support":
                support,
            "attack":
                attack,
            "reason":
                normalize_text(
                    item.get(
                        "reason"
                    )
                ),
        }

    for atom_id in atom_ids:

        if atom_id not in result:
            result[atom_id] = {
                "atom_id":
                    atom_id,
                "support":
                    [],
                "attack":
                    [],
                "reason":
                    "No validated "
                    "evidence alignment.",
            }

    return result


# ============================================================
# Deterministic four-valued logic
# ============================================================

def atom_state(
    alignment: dict,
) -> tuple[bool, bool]:

    return (
        bool(
            alignment[
                "support"
            ]
        ),
        bool(
            alignment[
                "attack"
            ]
        ),
    )


def state_name(
    state: tuple[bool, bool],
) -> str:

    true_support, false_support = (
        state
    )

    if (
        true_support
        and not false_support
    ):
        return "TRUE"

    if (
        false_support
        and not true_support
    ):
        return "FALSE"

    if (
        true_support
        and false_support
    ):
        return "BOTH"

    return "UNKNOWN"


def evaluate_expression(
    expression: dict,
    alignments: dict[str, dict],
) -> tuple[bool, bool]:

    op = expression["op"]

    if op == "CONST":

        if expression["value"]:
            return True, False

        return False, True

    if op == "ATOM":

        return atom_state(
            alignments[
                expression[
                    "atom_id"
                ]
            ]
        )

    children = [
        evaluate_expression(
            child,
            alignments,
        )
        for child
        in expression[
            "children"
        ]
    ]

    if op == "NOT":

        t, f = children[0]

        return f, t

    if op == "ALL":

        return (
            all(
                t
                for t, _
                in children
            ),
            any(
                f
                for _, f
                in children
            ),
        )

    if op == "ANY":

        return (
            any(
                t
                for t, _
                in children
            ),
            all(
                f
                for _, f
                in children
            ),
        )

    raise ValueError(
        f"Unknown operation {op}"
    )


def derive_outcome(
    applicability:
        tuple[bool, bool],
    satisfaction:
        tuple[bool, bool],
    non_applicability:
        tuple[bool, bool],
) -> str:

    app_t, app_f = applicability
    sat_t, sat_f = satisfaction
    na_t, na_f = non_applicability

    # Positive evidence that it applies
    # and positive evidence that it does not.
    if app_t and na_t:
        return "CONFLICTING"

    # N/A requires positive evidence.
    if na_t and not app_t:
        return "NOT_APPLICABLE"

    # Absence of applicability evidence
    # does NOT automatically mean N/A.
    if not app_t:
        return "UNKNOWN"

    if sat_t and sat_f:
        return "CONFLICTING"

    if sat_t:
        return "SATISFIED"

    if sat_f:
        return "VIOLATED"

    return "UNKNOWN"


def fold_submission_label(
    outcome: str,
) -> str | None:

    return {
        "SATISFIED":
            "1",
        "VIOLATED":
            "0",
        "NOT_APPLICABLE":
            "N/A",
    }.get(outcome)


# ============================================================
# Evaluate one case
# ============================================================

def evaluate_case(
    contract_path: Path,
    case_name: str,
    evidence_top_k: int,
) -> Path:

    print("\n" + "=" * 72)
    print("FRECA CORE V1 — EVALUATE")
    print("=" * 72)

    bundle = load_json(
        contract_path
    )

    cp = bundle["cp"]
    contract = bundle["contract"]

    print(
        f"\nCP: {cp['cp_id']}"
    )

    print(
        f"Contract model: "
        f"{bundle['model']}"
    )

    case_dir = find_case_dir(
        case_name
    )

    print(
        f"\nCase directory:\n"
        f"{case_dir}"
    )

    print(
        "\n[1/4] Parsing evidence..."
    )

    all_evidence = (
        load_case_evidence(
            case_dir
        )
    )

    print(
        f"Evidence chunks: "
        f"{len(all_evidence)}"
    )

    # --------------------------------------------------------
    # Requirement-level evidence reasoning pilot
    # D2.8 + D7.1 + D7.8 + D7.14
    # --------------------------------------------------------
    from evidence_reasoning_v2 import (
        run_from_evaluate_locals,
        print_requirement_result,
    )

    requirement_reasoning = run_from_evaluate_locals(
        locals(),
        retrieval_top_k=12,
    )

    print_requirement_result(
        requirement_reasoning
    )

    # --------------------------------------------------------
    # FRECA LEGACY FIREWALL — REQUIREMENT PIPELINE ONLY
    # V2 stops here. Legacy V1 atom/final-label logic remains
    # available when evaluate_case is called without this flag.
    # --------------------------------------------------------
    if globals().get(
        "FRECA_REQUIREMENT_PIPELINE_ONLY",
        False,
    ):
        print(
            "\n[LEGACY FIREWALL] "
            "Requirement pipeline complete; "
            "skipping legacy V1 atom/final-label tail."
        )
        return requirement_reasoning



    if not all_evidence:
        raise RuntimeError(
            "No evidence extracted."
        )

    query_parts = [
        cp["criterion"]
    ]

    query_parts.extend(
        atom["proposition"]
        for atom
        in contract["atoms"]
    )

    query = " ".join(
        query_parts
    )

    print(
        "\n[2/4] BM25 evidence "
        "retrieval..."
    )

    retrieved = bm25_rank(
        query,
        all_evidence,
        evidence_top_k,
    )

    for chunk in retrieved[:15]:

        print(
            f"  {chunk['id']}: "
            f"{chunk['score']:.3f}"
        )

    print(
        "\n[3/4] Evidence alignment "
        f"with {ALIGNMENT_MODEL}..."
    )

    raw_alignment = deepseek_json(
        model=ALIGNMENT_MODEL,
        system_prompt=
            ALIGNMENT_SYSTEM,
        user_prompt=
            make_alignment_prompt(
                contract,
                retrieved,
            ),
        thinking=False,
        max_tokens=7000,
    )

    alignments = validate_alignments(
        raw_alignment,
        contract["atoms"],
        retrieved,
    )

    atom_lookup = {
        atom["atom_id"]:
            atom
        for atom
        in contract["atoms"]
    }

    print(
        "\nValidated atom states:"
    )

    for atom_id in sorted(
        atom_lookup
    ):

        alignment = (
            alignments[
                atom_id
            ]
        )

        state = state_name(
            atom_state(
                alignment
            )
        )

        print(
            "\n"
            f"{atom_id} = {state}"
        )

        print(
            "  proposition:",
            atom_lookup[
                atom_id
            ][
                "proposition"
            ],
        )

        for item in alignment[
            "support"
        ]:

            print(
                "  SUPPORT:",
                f"[{item['chunk_id']}]",
                item["quote"],
            )

        for item in alignment[
            "attack"
        ]:

            print(
                "  ATTACK:",
                f"[{item['chunk_id']}]",
                item["quote"],
            )

    print(
        "\n[4/4] Deterministic "
        "evaluation..."
    )

    app_state = evaluate_expression(
        contract[
            "applicability"
        ],
        alignments,
    )

    sat_state = evaluate_expression(
        contract[
            "satisfaction"
        ],
        alignments,
    )

    na_state = evaluate_expression(
        contract[
            "non_applicability"
        ],
        alignments,
    )

    outcome = derive_outcome(
        app_state,
        sat_state,
        na_state,
    )

    label = fold_submission_label(
        outcome
    )

    print("\n" + "-" * 72)

    print(
        "Applicability      :",
        state_name(app_state),
    )

    print(
        "Satisfaction       :",
        state_name(sat_state),
    )

    print(
        "Non-applicability  :",
        state_name(na_state),
    )

    print(
        "Internal outcome   :",
        outcome,
    )

    print(
        "Submission label   :",
        (
            label
            if label is not None
            else "UNRESOLVED"
        ),
    )

    print("-" * 72)

    output_path = (
        RESULT_DIR
        / (
            f"{case_name}_"
            f"{cp['cp_id']}.json"
        )
    )

    save_json(
        {
            "schema":
                "freca-core-result-v1",
            "case":
                case_name,
            "case_dir":
                str(case_dir),
            "cp":
                cp,
            "contract_file":
                str(contract_path),
            "contract_model":
                bundle["model"],
            "alignment_model":
                ALIGNMENT_MODEL,
            "retrieved_evidence":
                retrieved,
            "alignments":
                alignments,
            "evaluation":
                {
                    "applicability":
                        state_name(
                            app_state
                        ),
                    "satisfaction":
                        state_name(
                            sat_state
                        ),
                    "non_applicability":
                        state_name(
                            na_state
                        ),
                    "internal_outcome":
                        outcome,
                    "submission_label":
                        label,
                },
        },
        output_path,
    )

    print(
        f"\nSaved:\n"
        f"{output_path}"
    )

    return output_path


# ============================================================
# CLI
# ============================================================

def show_cps():

    cps = load_official_cps()

    print(
        f"Found {len(cps)} CPs\n"
    )

    for cp in cps:

        print(
            f"{cp['cp_id']}: "
            f"{cp['criterion']}"
        )


def main():

    parser = argparse.ArgumentParser(
        description=
            "FRECA Core V1 pilot"
    )

    sub = parser.add_subparsers(
        dest="command",
        required=True,
    )

    sub.add_parser(
        "show-cps"
    )

    compile_parser = (
        sub.add_parser(
            "compile"
        )
    )

    compile_parser.add_argument(
        "--cp",
        required=True,
    )

    compile_parser.add_argument(
        "--policy-topk",
        type=int,
        default=30,
    )

    evaluate_parser = (
        sub.add_parser(
            "evaluate"
        )
    )

    evaluate_parser.add_argument(
        "--cp",
        required=True,
        help=
            "CP ID whose frozen "
            "contract should be used",
    )

    evaluate_parser.add_argument(
        "--case",
        required=True,
        help=
            "Case directory name, "
            "e.g. RE-NSW-2020-0144",
    )

    evaluate_parser.add_argument(
        "--evidence-topk",
        type=int,
        default=60,
    )

    args = parser.parse_args()

    CONTRACT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    RESULT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    if args.command == "show-cps":

        show_cps()
        return

    if args.command == "compile":

        compile_cp(
            args.cp,
            args.policy_topk,
        )
        return

    if args.command == "evaluate":

        cp_id = (
            canonical_cp_id(
                args.cp
            )
        )

        # Recover exact official spelling,
        # e.g. CP12
        cp = get_cp(cp_id)

        contract_path = (
            CONTRACT_DIR
            / f"{cp['cp_id']}.json"
        )

        if not contract_path.exists():

            raise FileNotFoundError(
                f"No frozen contract at "
                f"{contract_path}\n"
                f"Run:\n"
                f"python freca_core_v1.py "
                f"compile --cp "
                f"{cp['cp_id']}"
            )

        evaluate_case(
            contract_path,
            args.case,
            args.evidence_topk,
        )


if __name__ == "__main__":
    main()
